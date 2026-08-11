"""
Sherlock Reports & Exports Module
Handles saving query results into TXT, DOCX, and custom PDF formats.
"""

import os
import datetime
from typing import Dict, List, Any
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

class ReportGenerator:
    @staticmethod
    def export_txt(filepath: str, username: str, results: Dict[str, Any], phone_meta: Dict[str, Any] = None, phone_results: Dict[str, Any] = None):
        """
        Exports search results to a plain text (.txt) file.
        """
        with open(filepath, "w", encoding="utf-8") as f:
            f.write("=" * 60 + "\n")
            f.write("                 SHERLOCK OSINT SEARCH REPORT\n")
            f.write("=" * 60 + "\n\n")

            if phone_meta:
                f.write("[ TELEFOONNUMMER INFORMATIE ]\n")
                f.write(f"Ingevoerd nummer: {phone_meta.get('clean')}\n")
                f.write(f"E.164 indeling:   {phone_meta.get('e164')}\n")
                f.write(f"Internationaal:   {phone_meta.get('international')}\n")
                f.write(f"Nationaal:        {phone_meta.get('national')}\n")
                f.write(f"Type lijn:        {phone_meta.get('type')}\n")
                f.write(f"Provider:         {phone_meta.get('carrier')}\n")
                f.write(f"Locatie:          {phone_meta.get('location')}\n")
                f.write(f"Tijdzones:        {', '.join(phone_meta.get('timezones', []))}\n")
                f.write(f"Status:           {'Geldig nummer' if phone_meta.get('valid') else 'Ongeldig nummer'}\n\n")

                if phone_results:
                    for category, items in phone_results.items():
                        f.write(f"[ {category.upper()} ]\n")
                        if not items:
                            f.write("Geen vermeldingen gevonden op internet.\n\n")
                        else:
                            for i, item in enumerate(items, 1):
                                f.write(f" {i}. Titel:   {item.get('title')}\n")
                                f.write(f"    Link:    {item.get('url')}\n")
                                f.write(f"    Snippet: {item.get('snippet')}\n\n")
            else:
                f.write(f"Doelgebruikersnaam: {username}\n\n")
                f.write("[ GEPASTELDE ACCOUNTS GEPAST ]\n")
                claimed_count = 0
                for site, info in results.items():
                    # check if the result status is claimed
                    status_obj = info.get("status")
                    # handle both QueryResult object or dictionary
                    status_str = ""
                    if hasattr(status_obj, "status"):
                        status_str = str(status_obj.status)
                    else:
                        status_str = str(status_obj)

                    if "claimed" in status_str.lower() or "exists" in status_str.lower():
                        url_user = info.get("url_user", "")
                        f.write(f"- {site}: {url_user}\n")
                        claimed_count += 1

                f.write(f"\nTotaal gevonden accounts: {claimed_count}\n")

    @staticmethod
    def export_docx(filepath: str, username: str, results: Dict[str, Any], phone_meta: Dict[str, Any] = None, phone_results: Dict[str, Any] = None):
        """
        Exports search results to a Microsoft Word (.docx) file with structured scanning metadata.
        """
        doc = Document()
        doc.add_heading("Sherlock OSINT Zoekrapport", level=0)

        now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Scan Metadata block
        doc.add_heading("Scan Metadata", level=1)
        meta_table = doc.add_table(rows=1, cols=2)
        meta_hdr = meta_table.rows[0].cells
        meta_hdr[0].text = "Eigenschap"
        meta_hdr[1].text = "Waarde"

        # Add Scan Date
        row = meta_table.add_row().cells
        row[0].text = "Scan Datum"
        row[1].text = now_str

        if phone_meta:
            # Count phone matches
            matches_count = 0
            if phone_results:
                for cat, items in phone_results.items():
                    matches_count += len(items)

            row = meta_table.add_row().cells
            row[0].text = "Type Scan"
            row[1].text = "Telefoonnummer OSINT"

            row = meta_table.add_row().cells
            row[0].text = "Aantal online vermeldingen"
            row[1].text = str(matches_count)

            row = meta_table.add_row().cells
            row[0].text = "Status"
            row[1].text = "Succesvol voltooid"

            doc.add_heading("Telefoonnummer Informatie", level=1)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Eigenschap"
            hdr_cells[1].text = "Waarde"

            meta_keys = [
                ("Ingevoerd nummer", phone_meta.get("clean")),
                ("E.164 indeling", phone_meta.get("e164")),
                ("Internationaal", phone_meta.get("international")),
                ("Nationaal", phone_meta.get("national")),
                ("Type lijn", phone_meta.get("type")),
                ("Provider", phone_meta.get("carrier")),
                ("Locatie", phone_meta.get("location")),
                ("Tijdzones", ", ".join(phone_meta.get("timezones", []))),
                ("Validiteit", "Geldig" if phone_meta.get("valid") else "Ongeldig")
            ]

            for key, val in meta_keys:
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(val)

            if phone_results:
                for category, items in phone_results.items():
                    doc.add_heading(category, level=1)
                    if not items:
                        doc.add_paragraph("Geen online vermeldingen gevonden.")
                    else:
                        for item in items:
                            p = doc.add_paragraph()
                            p.add_run(f"Titel: {item.get('title')}\n").bold = True
                            p.add_run(f"Link: {item.get('url')}\n")
                            p.add_run(f"Snippet: {item.get('snippet')}\n")
        else:
            # Count matches and scan duration
            claimed_count = 0
            max_duration = 0.0
            for site, info in results.items():
                status_obj = info.get("status")
                status_str = ""
                if hasattr(status_obj, "status"):
                    status_str = str(status_obj.status)
                else:
                    status_str = str(status_obj)

                if "claimed" in status_str.lower() or "exists" in status_str.lower():
                    claimed_count += 1

                if hasattr(status_obj, "query_time") and status_obj.query_time is not None:
                    if status_obj.query_time > max_duration:
                        max_duration = status_obj.query_time

            duration_str = f"{max_duration:.2f} seconden" if max_duration > 0 else "Onbekend"

            row = meta_table.add_row().cells
            row[0].text = "Type Scan"
            row[1].text = "Gebruikersnaam OSINT"

            row = meta_table.add_row().cells
            row[0].text = "Aantal gevonden accounts"
            row[1].text = str(claimed_count)

            row = meta_table.add_row().cells
            row[0].text = "Totale scantijd"
            row[1].text = duration_str

            doc.add_paragraph(f"\nZoekopdracht uitgevoerd voor de gebruikersnaam: {username}")
            doc.add_heading("Gevonden Social Media Accounts", level=1)

            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Platform"
            hdr_cells[1].text = "Profiel URL"

            for site, info in results.items():
                status_obj = info.get("status")
                status_str = ""
                if hasattr(status_obj, "status"):
                    status_str = str(status_obj.status)
                else:
                    status_str = str(status_obj)

                if "claimed" in status_str.lower() or "exists" in status_str.lower():
                    row_cells = table.add_row().cells
                    row_cells[0].text = site
                    row_cells[1].text = info.get("url_user", "")

            doc.add_paragraph(f"\nTotaal aantal gevonden accounts: {claimed_count}")

        doc.save(filepath)

    @staticmethod
    def export_pdf(filepath: str, username: str, results: Dict[str, Any], phone_meta: Dict[str, Any] = None, phone_results: Dict[str, Any] = None):
        """
        Generates a beautifully styled, professional PDF report using ReportLab.
        """
        doc = SimpleDocTemplate(filepath, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
        story = []

        styles = getSampleStyleSheet()

        # Define modern professional styles
        title_style = ParagraphStyle(
            'ReportTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=26,
            textColor=colors.HexColor('#1A365D'), # Deep Navy Blue
            spaceAfter=15
        )

        subtitle_style = ParagraphStyle(
            'ReportSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=12,
            textColor=colors.HexColor('#4A5568'),
            spaceAfter=30
        )

        h1_style = ParagraphStyle(
            'SectionH1',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=16,
            textColor=colors.HexColor('#2B6CB0'), # Royal Blue
            spaceBefore=15,
            spaceAfter=10
        )

        h2_style = ParagraphStyle(
            'SectionH2',
            parent=styles['Heading3'],
            fontName='Helvetica-Bold',
            fontSize=12,
            textColor=colors.HexColor('#2D3748'),
            spaceBefore=10,
            spaceAfter=5
        )

        body_style = ParagraphStyle(
            'ReportBody',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=10,
            textColor=colors.HexColor('#2D3748'),
            leading=14
        )

        bold_body_style = ParagraphStyle(
            'ReportBodyBold',
            parent=body_style,
            fontName='Helvetica-Bold'
        )

        link_style = ParagraphStyle(
            'ReportLink',
            parent=body_style,
            textColor=colors.HexColor('#3182CE')
        )

        # Header Title
        story.append(Paragraph("Sherlock OSINT Onderzoeksrapport", title_style))
        if phone_meta:
            story.append(Paragraph(f"Doelwit Telefoonnummer: {phone_meta.get('international')}", subtitle_style))
        else:
            story.append(Paragraph(f"Doelwit Gebruikersnaam: {username}", subtitle_style))

        story.append(Spacer(1, 10))

        now_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Scan Metadata block
        story.append(Paragraph("Scan Metadata", h1_style))

        if phone_meta:
            matches_count = 0
            if phone_results:
                for cat, items in phone_results.items():
                    matches_count += len(items)

            scan_meta_data = [
                [Paragraph("Eigenschap", bold_body_style), Paragraph("Waarde", bold_body_style)],
                [Paragraph("Scan Datum", body_style), Paragraph(now_str, body_style)],
                [Paragraph("Type Scan", body_style), Paragraph("Telefoonnummer OSINT", body_style)],
                [Paragraph("Aantal online vermeldingen", body_style), Paragraph(str(matches_count), body_style)],
                [Paragraph("Status", body_style), Paragraph("Succesvol voltooid", body_style)]
            ]
        else:
            claimed_count = 0
            max_duration = 0.0
            for site, info in results.items():
                status_obj = info.get("status")
                status_str = ""
                if hasattr(status_obj, "status"):
                    status_str = str(status_obj.status)
                else:
                    status_str = str(status_obj)

                if "claimed" in status_str.lower() or "exists" in status_str.lower():
                    claimed_count += 1

                if hasattr(status_obj, "query_time") and status_obj.query_time is not None:
                    if status_obj.query_time > max_duration:
                        max_duration = status_obj.query_time

            duration_str = f"{max_duration:.2f} seconden" if max_duration > 0 else "Onbekend"

            scan_meta_data = [
                [Paragraph("Eigenschap", bold_body_style), Paragraph("Waarde", bold_body_style)],
                [Paragraph("Scan Datum", body_style), Paragraph(now_str, body_style)],
                [Paragraph("Type Scan", body_style), Paragraph("Gebruikersnaam OSINT", body_style)],
                [Paragraph("Aantal gevonden accounts", body_style), Paragraph(str(claimed_count), body_style)],
                [Paragraph("Totale scantijd", body_style), Paragraph(duration_str, body_style)]
            ]

        t_meta = Table(scan_meta_data, colWidths=[2.5 * inch, 4.5 * inch])
        t_meta.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (1,0), colors.HexColor('#E2E8F0')),
            ('TEXTCOLOR', (0,0), (-1,-1), colors.HexColor('#2D3748')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ('TOPPADDING', (0,0), (-1,-1), 5),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E0')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F7FAFC')]),
        ]))
        story.append(t_meta)
        story.append(Spacer(1, 20))

        if phone_meta:
            story.append(Paragraph("Telefoonnummer Validatie & Metadata", h1_style))

            # Metadata Table
            meta_data = [
                [Paragraph("Eigenschap", bold_body_style), Paragraph("Waarde", bold_body_style)],
                [Paragraph("Ingevoerd nummer", body_style), Paragraph(phone_meta.get("clean", ""), body_style)],
                [Paragraph("E.164 Formaat", body_style), Paragraph(phone_meta.get("e164", ""), body_style)],
                [Paragraph("Internationaal Formaat", body_style), Paragraph(phone_meta.get("international", ""), body_style)],
                [Paragraph("Nationaal Formaat", body_style), Paragraph(phone_meta.get("national", ""), body_style)],
                [Paragraph("Type Lijn", body_style), Paragraph(phone_meta.get("type", ""), body_style)],
                [Paragraph("Telecom Provider", body_style), Paragraph(phone_meta.get("carrier", ""), body_style)],
                [Paragraph("Geregistreerde Locatie", body_style), Paragraph(phone_meta.get("location", ""), body_style)],
                [Paragraph("Tijdzones", body_style), Paragraph(", ".join(phone_meta.get("timezones", [])), body_style)],
                [Paragraph("Status van nummer", body_style), Paragraph("Geldig en Actief" if phone_meta.get("valid") else "Ongeldig", body_style)]
            ]

            t = Table(meta_data, colWidths=[2.2 * inch, 4.8 * inch])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (1,0), colors.HexColor('#E2E8F0')),
                ('TEXTCOLOR', (0,0), (-1,-1), colors.HexColor('#2D3748')),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('BOTTOMPADDING', (0,0), (-1,0), 6),
                ('TOPPADDING', (0,0), (-1,0), 6),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E0')),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F7FAFC')]),
                ('TOPPADDING', (0,1), (-1,-1), 5),
                ('BOTTOMPADDING', (0,1), (-1,-1), 5),
            ]))
            story.append(t)
            story.append(Spacer(1, 20))

            if phone_results:
                story.append(PageBreak()) # Clean page break for search results
                story.append(Paragraph("Online Vermeldingen & OSINT Resultaten", h1_style))

                for category, items in phone_results.items():
                    story.append(Paragraph(category, h2_style))
                    story.append(Spacer(1, 5))

                    if not items:
                        story.append(Paragraph("Geen vermeldingen gevonden op openbare netwerken.", body_style))
                        story.append(Spacer(1, 10))
                    else:
                        for item in items:
                            # Print result blocks professionally
                            res_text = f"<b>Titel:</b> {item.get('title')}<br/>" \
                                       f"<b>Link:</b> <font color='#3182CE'>{item.get('url')}</font><br/>" \
                                       f"<b>Snippet:</b> {item.get('snippet')}"
                            story.append(Paragraph(res_text, body_style))
                            story.append(Spacer(1, 12))
        else:
            story.append(Paragraph("Gevonden Social Media Accounts", h1_style))

            # Accounts Table
            account_data = [
                [Paragraph("Social Media Netwerk", bold_body_style), Paragraph("Profiel Link / Gebruikersnaam URL", bold_body_style)]
            ]

            claimed_count = 0
            for site, info in results.items():
                status_obj = info.get("status")
                status_str = ""
                if hasattr(status_obj, "status"):
                    status_str = str(status_obj.status)
                else:
                    status_str = str(status_obj)

                if "claimed" in status_str.lower() or "exists" in status_str.lower():
                    url_user = info.get("url_user", "")
                    account_data.append([
                        Paragraph(site, body_style),
                        Paragraph(f"<a href='{url_user}'>{url_user}</a>", link_style)
                    ])
                    claimed_count += 1

            if claimed_count == 0:
                story.append(Paragraph("Er zijn geen accounts gevonden voor deze gebruikersnaam.", body_style))
            else:
                t = Table(account_data, colWidths=[2.5 * inch, 4.5 * inch])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (1,0), colors.HexColor('#3182CE')),
                    ('TEXTCOLOR', (0,0), (1,0), colors.white),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('BOTTOMPADDING', (0,0), (-1,0), 6),
                    ('TOPPADDING', (0,0), (-1,0), 6),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#E2E8F0')),
                    ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F7FAFC')]),
                    ('TOPPADDING', (0,1), (-1,-1), 5),
                    ('BOTTOMPADDING', (0,1), (-1,-1), 5),
                ]))
                story.append(t)
                story.append(Spacer(1, 20))
                story.append(Paragraph(f"<b>Totaal gevonden accounts:</b> {claimed_count}", body_style))

        doc.build(story)
